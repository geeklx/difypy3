import json
import os
import sys
# 将项目根目录添加到 Python 路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import subprocess
import time
from datetime import datetime
import logging
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi.staticfiles import StaticFiles
from fastapi import FastAPI, HTTPException, APIRouter, Request, Form, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.responses import JSONResponse
from openai import OpenAI
from geekaiapp.g_model import TTSRequest, VideoSubmission, AudioSubmission, VideoRequest, VideoResponse, JMRequest
from geekaiapp.g_utils import save_audio_file, output_path1, upload_cos, tencent_region, tencent_secret_id, \
    tencent_secret_key, tencent_bucket, siliconflow_api_url, siliconflow_auth_token, gjld_submit_video_job, \
    gjld_check_video_status, siliconflow_videomodel, siliconflow_audiomodel, siliconflow_voice, zpai_video_job, \
    zpai_check_video_status, microsoft_api_key, microsoft_base_url, aliyuncs_api_key, aliyuncs_base_url, ip, ip_tts, \
    aliyuncs_model, generate_clip, marp_path, ip_md, ip_html, port



app = FastAPI()
router = APIRouter()



# 获取当前文件的目录
current_dir = os.path.dirname(os.path.abspath(__file__))
# 拼接 static 目录的绝对路径
static_dir = os.path.join(current_dir, "static")
# 挂载静态文件目录
router.mount("/static", StaticFiles(directory=static_dir), name="static")

# 挂载静态文件目录
# router.mount("/static", StaticFiles(directory="static"), name="static")

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 微软服务
client = OpenAI(
    api_key=microsoft_api_key,
    base_url=microsoft_base_url
)

# 阿里大模型服务
client2 = OpenAI(
    api_key=aliyuncs_api_key,
    base_url=aliyuncs_base_url
)


# g微软的文生音1
@router.post('/edge/tts1/')
async def generate_tts1(request_body: TTSRequest):
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        filename, output_path = save_audio_file(response.content, output_path1)
        audio_url = f"{ip}{ip_tts}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url
        }
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# g微软的文生音，并上传到腾讯oss
@router.post("/edge/tts12/")
async def generate_tts12(request_body: TTSRequest):
    """
    Generates text-to-speech audio using the edge tts API and uploads it to Tencent Cloud COS.
    """
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        # Save the audio file
        filename, output_path2 = save_audio_file(response.content, output_path1)
        audio_url = f"{ip}{ip_tts}/{filename}"
        # Upload to COS
        etag = upload_cos('text1', tencent_region, tencent_secret_id, tencent_secret_key, tencent_bucket, filename,
                          output_path1)
        if etag:
            audio_url2 = f"https://{tencent_bucket}.cos.{tencent_region}.myqcloud.com/{filename}"
            return {
                "audio_url": audio_url2,
                "filename": filename,
                "output_path": audio_url,
                "etag": etag
            }
        else:
            raise HTTPException(status_code=500, detail="Failed to upload audio to COS")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 硅基流动-文生视频
@router.post("/gjld/video/")
async def sjld_video(video_submission: VideoSubmission):
    # Use the model from the user input if provided, otherwise use the default from config
    model = video_submission.model if video_submission.model else siliconflow_videomodel

    submit_response = gjld_submit_video_job(siliconflow_api_url, siliconflow_auth_token, model, video_submission.prompt)
    if "error" in submit_response:
        raise HTTPException(status_code=500, detail=submit_response["error"])

    request_id = submit_response.get("requestId")
    if not request_id:
        raise HTTPException(status_code=500, detail="Failed to get requestId from submit response.")

    status_response = gjld_check_video_status(siliconflow_api_url, siliconflow_auth_token, request_id)
    if status_response and status_response.get("status") == "Succeed" and status_response.get("results") and \
            status_response["results"].get("videos") and len(status_response["results"]["videos"]) > 0:
        video_url = status_response['results']['videos'][0]['url']
        return {"video_url": video_url}
    else:
        raise HTTPException(status_code=500, detail="Invalid status response or missing video URL.")


# g硅基流动-文生音
@router.post("/gjld/audio/")
async def gjld_audio(audio_Submission: AudioSubmission):
    audiourl = siliconflow_api_url + "/audio/speech"
    audiomodel2 = audio_Submission.model if audio_Submission.model else siliconflow_audiomodel
    voice2 = audio_Submission.voice if audio_Submission.voice else siliconflow_voice
    headers = {
        "Authorization": f"Bearer {siliconflow_auth_token}",
        "Content-Type": "application/json"
    }
    data = {
        "model": audiomodel2,
        "input": audio_Submission.input,
        "voice": voice2,
        "response_format": "mp3",
        "sample_rate": 32000,
        "stream": True,
        "speed": 1,
        "gain": 0
    }
    response = requests.post(audiourl, headers=headers, json=data)

    if response.status_code == 200:
        filename, output_path2 = save_audio_file(response.content, output_path1)
        audio_url = f"{ip}{ip_tts}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url
        }
        # env = 'test'
        # etag = upload_cos(env, filename, output_path)
        # return {
        #     "filename": filename,
        #     "output_path": output_path2,
        #     "etag": etag
        # }
    else:
        raise HTTPException(status_code=response.status_code, detail="请求失败")


# g智谱AI-文生视频任务的接口
@router.post("/zhipuai/video/")
async def zpai_video(video_request: VideoRequest):
    """Submits a text-to-video generation job using the ZhipuAI API."""
    try:
        task_id = zpai_video_job(video_request.prompt, video_request.with_audio)
        logger.info(f"Video generation task submitted successfully. Task ID: {task_id}")

        # 检查任务状态并返回结果
        status_response = zpai_check_video_status(task_id)
        return VideoResponse(**status_response)

    except HTTPException as err:
        raise err
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")


# g即梦-文生图的接口 https://github.com/LLM-Red-Team/jimeng-free-api
@router.post("/jimeng/img/")
async def jm_image(request1: JMRequest):
    headers = {
        "Authorization": f"Bearer {request1.image_api_key}",
        "Content-Type": "application/json",
    }
    data = {
        "model": request1.model,
        "prompt": request1.prompt,
        "negativePrompt": request1.negativePrompt,
        "width": request1.width,
        "height": request1.height,
        "sample_strength": request1.sample_strength,
    }
    try:
        response = requests.post(request1.image_generation_url, headers=headers, json=data)
        response.raise_for_status()  # 检查 HTTP 状态码
        result = response.json()
        return {
            "url": result["data"][0]["url"]
        }
    except requests.exceptions.RequestException as e:
        raise Exception(f"图像生成 API 请求失败: {e}")
    except (KeyError, IndexError) as e:
        raise Exception(f"图像生成 API 响应解析失败: {e}")


system_prompt = """
你是一位英语教学专家，请根据用户将提供给你的内容，请你分析内容，并提取其中的关键信息，识别出中文对应的英文写错的部分,以 JSON 的形式输出，输出的 JSON 需遵守以下的格式：

示例输入:
{
    "content": [
        {
            "序号": 1,
            "汉语": "n.哨兵 n./vt.守卫,保卫,看守",
            "英语": "guard"
        },
        {
            "序号": 2,
            "汉语": "vt.保护,护卫 n.安全装置",
            "英语": "safeguard"
        }
    ]
}

JSON 输出示例:
{
  "errors": [
    {
      "序号": <1>,
      "汉语": <n.保证(书),保修单 vt.保证,提供(产品)保修单>,
      "错误英语": <gooe>,
      "正确英语": <gruarantee>
    }
  ]
}
"""


# g单词比对
@router.post("/dcbd1/")
async def get_dps123(request: Request):
    try:
        data = await request.json()
        # print(f'data ={data}')
        text = data.get('content')
        # print(f'text ={text}')
        response = client2.chat.completions.create(
            model=aliyuncs_model,
            # messages=messages,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": json.dumps(text, ensure_ascii=False)
                }
            ],
            response_format={
                'type': 'json_object'
            }
        )
        # print(json.loads(response.choices[0].message.content))
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# bizy绘画-commfyui
@router.post("/comfyui_bizyairapi/")
# async def generate_clip_endpoint(request: GenerateClipRequest,file: UploadFile = File(...)):
async def generate_clip_endpoint(prompt: str = Form(...), seed: int = Form(...), idx: int = Form(...),
                                 workflowfile: UploadFile = File(...)):
    try:
        # logger.info(f"Received request: {request}")
        logger.info(f"Received request: prompt={prompt}, seed={seed}, idx={idx}, workflowfile={workflowfile.filename}")
        # 读取上传的文件内容
        workflow_content = await workflowfile.read()
        # 将文件内容转换为json
        workflow_data = json.loads(workflow_content)
        filename, output_path, etag = generate_clip(
            prompt,
            seed,
            workflow_data,
            idx
        )
        return {
            "filename": filename,
            "output_path": output_path,
            "etag": etag
        }
    except Exception as e:
        logger.error(f"Error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# x保存上传的Markdown内容，并转换成PPT
@router.post('/pptupload/')
async def upload_markdown(request: Request):
    content = await request.body()  # 异步获取请求体
    content = content.decode('utf-8')  # 将字节转换为字符串
    timestamp = str(int(time.time()))
    md_filename = f"{timestamp}.md"
    pptx_filename = f"{timestamp}.pptx"
    # 保存Markdown文件
    with open(f"{ip_md}/{md_filename}", 'w', encoding='utf-8') as f:
        f.write(content)
    # 使用marp-cli将Markdown转换为PPT
    try:
        subprocess.run([marp_path, f'{ip_md}/{md_filename}', '-o', f'{ip_html}/{pptx_filename}'], shell=True,
                       check=True)
    except subprocess.CalledProcessError as e:
        return {
            'message': 'Failed to convert Markdown to PPT',
            'error': str(e)
        }
    # 返回文件链接
    audio_url = f"{ip}{ip_md}/{pptx_filename}"
    return f'Markdown 文件已保存\n预览链接: {ip}{ip_md}/{md_filename} \n下载链接: {ip}{ip_html}/{pptx_filename}?pptx'


# g获取网页信息到word
@router.post('/generate_doc/')
async def generate_doc(request: Request):
    try:
        # 获取请求中的JSON数据
        data = await request.json()
        title = data.get('title')
        content = data.get('content')

        if not title and not content:
            logger.error("Title or content is required")
            return JSONResponse({"error": "Title or content is required"}, status_code=400)

        # 生成文档
        file_name = f"llm_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(ip_html, file_name)
        logger.debug(f"File path: {file_path}")

        doc = Document()

        if title:
            # 添加大标题
            paragraph = doc.add_heading(title, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(22)  # 二号字体

        if content:
            # 添加正文
            paragraph = doc.add_paragraph(content)
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(10.5)  # 五号字体

        doc.save(file_path)
        logger.info(f"Document generated successfully at {file_path}")

        # 在Mac上打开文件
        # subprocess.call(['open', file_path], shell=True)
        word_url = f"{ip}{ip_html}/{file_name}"
        return JSONResponse({"message": "Document generated successfully", "file_path": word_url}, status_code=200)

    except Exception as e:
        logger.error(f"Error generating document: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# g上传并处理Markdown文件的路径
@router.post('/markdown2map/upload/')
async def upload_markdown2map(request: Request):
    content = await request.body()
    content = content.decode('utf-8')
    time_name = str(int(time.time()))  # 生成时间戳作为文件名
    md_file_name = time_name + ".md"  # Markdown文件名
    html_file_name = time_name + ".html"  # HTML文件名

    # 创建markdown和html文件夹，如果它们不存在的话
    os.makedirs('static/markdown', exist_ok=True)
    os.makedirs('static/html', exist_ok=True)

    # 将Markdown内容写入文件
    with open(f'{ip_md}/{md_file_name}', "w", encoding='utf-8') as f:
        f.write(content)

    print(f"Markdown file created: {ip_md}/{md_file_name}")

    # 使用subprocess调用markmap-cli将Markdown转换为HTML，并移动到static/html目录
    try:
        result = subprocess.run(['npx', 'markmap-cli', f'{ip_md}/{md_file_name}', '--no-open'], capture_output=True,
                                shell=True,
                                text=True)

        if result.returncode != 0:
            raise subprocess.CalledProcessError(result.returncode, result.args, output=result.stdout,
                                                stderr=result.stderr)

        # 尝试将生成的HTML文件移动到static/html文件夹
        os.replace(f'{ip_md}/{html_file_name}', f'{ip_html}/{html_file_name}')
        print(f"HTML file moved to: {ip_html}/{html_file_name}")

        # 返回转换后的HTML文件链接
        # return f'Markdown文件已保存. 点击预览: {request.url_for("get_html", filename=html_file_name)}'
        return f'Markdown文件已保存. 点击预览: {ip}{ip_html}/{html_file_name}'
    except subprocess.CalledProcessError as e:
        # 如果转换过程中出现错误，返回错误信息
        return f"Error generating HTML file: {e.output}\n{e.stderr}", 500


# g提供HTML文件的路径
@router.post('/static/html/{filename}')
async def get_html(filename: str):
    return FileResponse(f'static/html/{filename}')




import uvicorn
from g_jiekou import router as router_edgetts

app.include_router(router_edgetts)

if __name__ == '__main__':
    uvicorn.run(router, host='0.0.0.0', port=port)
