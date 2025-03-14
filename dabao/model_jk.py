import json
import os
import subprocess
import time
from datetime import datetime

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import FastAPI, HTTPException, APIRouter, Request, Form, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI

from dabao.cos_model import TTSRequest, VideoSubmission, AudioSubmission, VideoRequest, VideoResponse, JMRequest
# sys.path.append('path/to/dabao')  # 添加模块路径到搜索路径
from dabao.cos_utils import openai_api_key1, openai_base_url1, port1, ip1, region1, secret_id1, secret_key1, bucket1, \
    save_audio_file, upload_cos, videomodel, siliconflow_api_url, siliconflow_auth_token, \
    audiomodel, voice, logger, output_path1, gjld_check_video_status, zpai_check_video_status, \
    api_key1, base_url1, model1, system_prompt, gjld_submit_video_job, zpai_video_job, generate_clip

app = FastAPI()

router = APIRouter()

# 基础配置
api_key = openai_api_key1
base_url = openai_base_url1
output_path = output_path1
port = port1
ip = ip1

# Tencent Cloud COS configuration
region = region1
secret_id = secret_id1
secret_key = secret_key1
bucket = bucket1

# 微软服务
client = OpenAI(
    api_key=api_key,
    base_url=base_url
)

# 阿里大模型服务
client2 = OpenAI(
    api_key=api_key1,
    base_url=base_url1
)


# 微软的文生音，并上传到腾讯oss
@router.post("/edge/tts12/")
async def generate_tts12(request_body: TTSRequest):
    """
    Generates text-to-speech audio using the edge tts API and uploads it to Tencent Cloud COS.
    """
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        # Save the audio file
        filename, output_path2 = save_audio_file(response.content, output_path)
        ip = 'http://192.168.2.2:5031/dabaotmp1'
        audio_url2 = f"{ip}/{filename}"
        # Upload to COS
        etag = upload_cos(region, secret_id, secret_key, bucket, filename, output_path)
        if etag:
            audio_url = f"https://{bucket}.cos.{region}.myqcloud.com/{filename}"
            return {
                "audio_url": audio_url,
                "filename": filename,
                "output_path": audio_url2,
                "etag": etag
            }
        else:
            raise HTTPException(status_code=500, detail="Failed to upload audio to COS")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 微软的文生音1
@router.post('/edge/tts1/')
async def generate_tts1(request_body: TTSRequest):
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        filename, output_path2 = save_audio_file(response.content, output_path)
        ip = 'http://192.168.31.116:5052/tts1'
        audio_url2 = f"{ip}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url2
        }
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 硅基流动-文生视频
@router.post("/gjld/video/")
async def sjld_video(video_submission: VideoSubmission):
    # Use the model from the user input if provided, otherwise use the default from config
    model = video_submission.model if video_submission.model else videomodel

    submit_response = gjld_submit_video_job(siliconflow_api_url, siliconflow_auth_token, model, video_submission.prompt)
    if "error" in submit_response:
        raise HTTPException(status_code=500, detail=submit_response["error"])

    request_id = submit_response.get("requestId")
    if not request_id:
        raise HTTPException(status_code=500, detail="Failed to get requestId from submit response.")

    status_response = gjld_check_video_status(siliconflow_api_url, siliconflow_auth_token, request_id)
    if status_response and status_response.get("status") == "Succeed" and status_response.get("results") and \
            status_response["results"].get("videos") and len(status_response["results"]["videos"]) > 0:
        video_url = status_response['results']['videos'][0]['url']
        return {"video_url": video_url}
    else:
        raise HTTPException(status_code=500, detail="Invalid status response or missing video URL.")


# 硅基流动-文生音
@router.post("/gjld/audio/")
async def gjld_audio(audio_Submission: AudioSubmission):
    audiourl = siliconflow_api_url + "/audio/speech"
    audiomodel2 = audio_Submission.model if audio_Submission.model else audiomodel
    voice2 = audio_Submission.voice if audio_Submission.voice else voice
    headers = {
        "Authorization": f"Bearer {siliconflow_auth_token}",
        "Content-Type": "application/json"
    }
    data = {
        "model": audiomodel2,
        "input": audio_Submission.input,
        "voice": voice2,
        "response_format": "mp3",
        "sample_rate": 32000,
        "stream": True,
        "speed": 1,
        "gain": 0
    }
    response = requests.post(audiourl, headers=headers, json=data)

    if response.status_code == 200:
        filename, output_path2 = save_audio_file(response.content, output_path)
        ip = 'http://10.20.0.65:5055/g2'
        audio_url2 = f"{ip}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url2
        }
        # env = 'test'
        # etag = upload_cos(env, filename, output_path)
        # return {
        #     "filename": filename,
        #     "output_path": output_path2,
        #     "etag": etag
        # }
    else:
        raise HTTPException(status_code=response.status_code, detail="请求失败")


# 智谱AI-文生视频任务的接口
@router.post("/zhipuai/video/")
async def zpai_video(video_request: VideoRequest):
    """Submits a text-to-video generation job using the ZhipuAI API."""
    try:
        task_id = zpai_video_job(video_request.prompt, video_request.with_audio)
        logger.info(f"Video generation task submitted successfully. Task ID: {task_id}")

        # 检查任务状态并返回结果
        status_response = zpai_check_video_status(task_id)
        return VideoResponse(**status_response)

    except HTTPException as err:
        raise err
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")


# 即梦-文生图的接口
@router.post("/jimeng/img/")
async def jm_image(request1: JMRequest):
    headers = {
        "Authorization": f"Bearer {request1.image_api_key}",
        "Content-Type": "application/json",
    }
    data = {
        "model": request1.model,
        "prompt": request1.prompt,
        "negativePrompt": request1.negativePrompt,
        "width": request1.width,
        "height": request1.height,
        "sample_strength": request1.sample_strength,
    }
    try:
        response = requests.post(request1.image_generation_url, headers=headers, json=data)
        response.raise_for_status()  # 检查 HTTP 状态码
        result = response.json()
        return {
            "url": result["data"][0]["url"]
        }
    except requests.exceptions.RequestException as e:
        raise Exception(f"图像生成 API 请求失败: {e}")
    except (KeyError, IndexError) as e:
        raise Exception(f"图像生成 API 响应解析失败: {e}")


# 单词比对
@router.post("/dcbd1/")
async def get_dps123(request: Request):
    try:
        data = await request.json()
        # print(f'data ={data}')
        text = data.get('content')
        # print(f'text ={text}')
        response = client2.chat.completions.create(
            model=model1,
            # messages=messages,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": json.dumps(text, ensure_ascii=False)
                }
            ],
            response_format={
                'type': 'json_object'
            }
        )
        # print(json.loads(response.choices[0].message.content))
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# bizy绘画-commfyui
@router.post("/comfyui_bizyairapi/")
# async def generate_clip_endpoint(request: GenerateClipRequest,file: UploadFile = File(...)):
async def generate_clip_endpoint(prompt: str = Form(...), seed: int = Form(...), idx: int = Form(...),
                                 workflowfile: UploadFile = File(...)):
    try:
        # logger.info(f"Received request: {request}")
        logger.info(f"Received request: prompt={prompt}, seed={seed}, idx={idx}, workflowfile={workflowfile.filename}")
        # 读取上传的文件内容
        workflow_content = await workflowfile.read()
        # 将文件内容转换为json
        workflow_data = json.loads(workflow_content)
        filename, output_path, etag = generate_clip(
            prompt,
            seed,
            workflow_data,
            idx
        )
        return {
            "filename": filename,
            "output_path": output_path,
            "etag": etag
        }
    except Exception as e:
        logger.error(f"Error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


marp_path = os.getenv('MARP_PATH')
# 检查是否获取到了路径
if marp_path:
    print(f"Marp 的路径是: {marp_path}")
else:
    print("未找到 Marp 的路径，请确保环境变量 MARP_PATH 已设置。")


# 保存上传的Markdown内容，并转换成PPT
@router.post('/pptupload/')
async def upload_markdown(request: Request):
    content = await request.body()  # 异步获取请求体
    content = content.decode('utf-8')  # 将字节转换为字符串
    timestamp = str(int(time.time()))
    md_filename = f"{timestamp}.md"
    pptx_filename = f"{timestamp}.pptx"
    # 保存Markdown文件
    with open(f"data/{md_filename}", 'w', encoding='utf-8') as f:
        f.write(content)
    # 使用marp-cli将Markdown转换为PPT
    try:
        subprocess.run([marp_path, f'tmp/{md_filename}', '-o', f'tmp/{pptx_filename}'], check=True)
    except subprocess.CalledProcessError as e:
        return {
            'message': 'Failed to convert Markdown to PPT',
            'error': str(e)
        }
    # 返回文件链接
    return f'Markdown 文件已保存\n预览链接: http://127.0.0.1:5037/tmp/{md_filename} \n下载链接: http://127.0.0.1:5037/tmp/{pptx_filename}?pptx'


# 配置保存文档的目录
SAVE_DIR = "./tmp/"
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)


# 获取网页信息到word
@router.post('/generate_doc/')
async def generate_doc(request: Request):
    try:
        # 获取请求中的JSON数据
        data = await request.json()
        title = data.get('title')
        content = data.get('content')

        if not title and not content:
            logger.error("Title or content is required")
            return JSONResponse({"error": "Title or content is required"}, status_code=400)

        # 生成文档
        file_name = f"llm_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(SAVE_DIR, file_name)
        logger.debug(f"File path: {file_path}")

        doc = Document()

        if title:
            # 添加大标题
            paragraph = doc.add_heading(title, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(22)  # 二号字体

        if content:
            # 添加正文
            paragraph = doc.add_paragraph(content)
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(10.5)  # 五号字体

        doc.save(file_path)
        logger.info(f"Document generated successfully at {file_path}")

        # 在Mac上打开文件
        subprocess.call(['open', file_path], shell=True)

        return JSONResponse({"message": "Document generated successfully", "file_path": file_path}, status_code=200)

    except Exception as e:
        logger.error(f"Error generating document: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# 上传并处理Markdown文件的路径
@router.post('/markdown2map/upload/')
async def upload_markdown2map(request: Request):
    content = await request.body()
    content = content.decode('utf-8')
    time_name = str(int(time.time()))  # 生成时间戳作为文件名
    md_file_name = time_name + ".md"  # Markdown文件名
    html_file_name = time_name + ".html"  # HTML文件名

    # 创建markdown和html文件夹，如果它们不存在的话
    os.makedirs('markdown', exist_ok=True)
    os.makedirs('static/html', exist_ok=True)

    # 将Markdown内容写入文件
    with open(f'markdown/{md_file_name}', "w", encoding='utf-8') as f:
        f.write(content)

    print(f"Markdown file created: markdown/{md_file_name}")

    # 使用subprocess调用markmap-cli将Markdown转换为HTML，并移动到static/html目录
    try:
        result = subprocess.run(['npx', 'markmap-cli', f'markdown/{md_file_name}', '--no-open'], capture_output=True,
                                shell=True,
                                text=True)

        if result.returncode != 0:
            raise subprocess.CalledProcessError(result.returncode, result.args, output=result.stdout,
                                                stderr=result.stderr)

        # 尝试将生成的HTML文件移动到static/html文件夹
        os.replace(f'markdown/{html_file_name}', f'static/html/{html_file_name}')
        print(f"HTML file moved to: static/html/{html_file_name}")

        # 返回转换后的HTML文件链接
        return f'Markdown文件已保存. 点击预览: {request.url_for("get_html", filename=html_file_name)}'
    except subprocess.CalledProcessError as e:
        # 如果转换过程中出现错误，返回错误信息
        return f"Error generating HTML file: {e.output}\n{e.stderr}", 500

# 挂载静态文件目录
app.mount("/static", StaticFiles(directory="static"), name="static")

# 提供HTML文件的路径
@router.post('/html/{filename}')
async def get_html(filename: str):
    return FileResponse(f'static/html/{filename}')

# import uvicorn
#
# if __name__ == '__main__':
#     uvicorn.run(app, host='0.0.0.0', port=5036)
